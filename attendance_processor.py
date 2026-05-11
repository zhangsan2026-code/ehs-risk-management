import os
import re
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

class AttendanceProcessor:
    def __init__(self):
        self.attendance_records = []
    
    def parse_excel_attendance(self, file_path):
        try:
            df = pd.read_excel(file_path)
            return self._extract_from_dataframe(df)
        except Exception as e:
            print(f"读取Excel文件失败: {e}")
            return None
    
    def parse_docx_attendance(self, file_path):
        doc = Document(file_path)
        data = []
        
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                if row_data and any(row_data):
                    data.append(row_data)
        
        return self._extract_from_list(data)
    
    def _extract_from_dataframe(self, df):
        result = {
            'title': '',
            'date': '',
            'time': '',
            'location': '',
            'department': '',
            'total_expected': 0,
            'total_attended': 0,
            'total_absent': 0,
            'attendees': [],
            'absentees': []
        }
        
        for col in df.columns:
            if '日期' in str(col) or '时间' in str(col):
                for val in df[col].dropna():
                    if str(val):
                        date_match = re.search(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日号]?)', str(val))
                        if date_match:
                            result['date'] = date_match.group(1)
                        time_match = re.search(r'(\d{1,2}:\d{2}.*)', str(val))
                        if time_match:
                            result['time'] = time_match.group(1)
        
        name_col = None
        status_col = None
        
        for col in df.columns:
            if any(keyword in str(col) for keyword in ['姓名', '人员', '员工', '参会人员']):
                name_col = col
            elif any(keyword in str(col) for keyword in ['签到', '状态', '出勤', '是否']):
                status_col = col
        
        if name_col:
            for _, row in df.iterrows():
                name = str(row[name_col]).strip()
                if not name or name in ['姓名', '人员', 'nan']:
                    continue
                
                status = '已签到'
                if status_col:
                    status_val = str(row[status_col]).strip()
                    if status_val in ['缺席', '未签到', '请假', '请假中', '否']:
                        status = '未签到'
                    elif status_val in ['已签到', '出勤', '是']:
                        status = '已签到'
                
                if status == '已签到':
                    result['attendees'].append({'name': name, 'status': '已签到'})
                    result['total_attended'] += 1
                else:
                    result['absentees'].append({'name': name, 'status': status})
                    result['total_absent'] += 1
        
        result['total_expected'] = result['total_attended'] + result['total_absent']
        return result
    
    def _extract_from_list(self, data):
        result = {
            'title': '',
            'date': '',
            'time': '',
            'location': '',
            'department': '',
            'total_expected': 0,
            'total_attended': 0,
            'total_absent': 0,
            'attendees': [],
            'absentees': []
        }
        
        for row in data:
            row_str = ' '.join(row)
            
            if '日期' in row_str or '时间' in row_str:
                date_match = re.search(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日号]?)', row_str)
                if date_match:
                    result['date'] = date_match.group(1)
                time_match = re.search(r'(\d{1,2}:\d{2}.*)', row_str)
                if time_match:
                    result['time'] = time_match.group(1)
            elif '部门' in row_str:
                result['department'] = re.sub(r'部门\s*[：:]', '', row_str).strip()
            elif '地点' in row_str:
                result['location'] = re.sub(r'地点\s*[：:]', '', row_str).strip()
            
            name_pattern = re.compile(r'([\u4e00-\u9fa5]{2,4})')
            names = name_pattern.findall(row_str)
            
            if names:
                status = '已签到'
                if any(word in row_str for word in ['缺席', '未签到', '请假']):
                    status = '未签到'
                
                for name in names:
                    if name not in ['姓名', '部门', '日期', '时间', '地点', '签到']:
                        if status == '已签到':
                            result['attendees'].append({'name': name, 'status': '已签到'})
                            result['total_attended'] += 1
                        else:
                            result['absentees'].append({'name': name, 'status': status})
                            result['total_absent'] += 1
        
        result['total_expected'] = result['total_attended'] + result['total_absent']
        return result
    
    def generate_report(self, data, output_path):
        doc = Document()
        
        style = doc.styles['Heading 1']
        style.font.size = Pt(16)
        style.font.bold = True
        
        doc.add_heading(f"{data['department'] or '培训/会议'}签到统计报告", level=1)
        
        info_table = doc.add_table(rows=2, cols=3)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = '日期'
        info_table.cell(0, 1).text = data['date'] or datetime.now().strftime('%Y年%m月%d日')
        info_table.cell(0, 2).text = '时间'
        info_table.cell(1, 0).text = data['time'] or '未指定'
        info_table.cell(1, 1).text = '地点'
        info_table.cell(1, 2).text = data['location'] or '未指定'
        
        for row in info_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        stats_table = doc.add_table(rows=1, cols=4)
        stats_table.style = 'Table Grid'
        stats_table.cell(0, 0).text = '应到人数'
        stats_table.cell(0, 1).text = str(data['total_expected'])
        stats_table.cell(0, 2).text = '实到人数'
        stats_table.cell(0, 3).text = str(data['total_attended'])
        
        rate = (data['total_attended'] / data['total_expected'] * 100) if data['total_expected'] > 0 else 0
        
        doc.add_paragraph(f"\n签到率: {rate:.1f}%")
        
        if data['attendees']:
            doc.add_heading('签到人员', level=2)
            attendees_table = doc.add_table(rows=1, cols=2)
            attendees_table.style = 'Table Grid'
            attendees_table.cell(0, 0).text = '序号'
            attendees_table.cell(0, 1).text = '姓名'
            
            for i, attendee in enumerate(data['attendees'], 1):
                row = attendees_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = attendee['name']
        
        if data['absentees']:
            doc.add_heading('缺席人员', level=2)
            absentees_table = doc.add_table(rows=1, cols=2)
            absentees_table.style = 'Table Grid'
            absentees_table.cell(0, 0).text = '序号'
            absentees_table.cell(0, 1).text = '姓名'
            
            for i, absentee in enumerate(data['absentees'], 1):
                row = absentees_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = absentee['name']
        
        doc.save(output_path)
        return output_path
    
    def batch_process(self, input_dir, output_dir):
        os.makedirs(output_dir, exist_ok=True)
        results = []
        
        for filename in os.listdir(input_dir):
            if filename.endswith('.xlsx') or filename.endswith('.xls'):
                if '签到' in filename or '考勤' in filename:
                    source_path = os.path.join(input_dir, filename)
                    data = self.parse_excel_attendance(source_path)
                    
                    if data:
                        date_str = data['date'].replace('/', '-').replace('年', '-').replace('月', '-').replace('日', '').replace('号', '')
                        date_str = re.sub(r'[-]+', '-', date_str).strip('-')
                        
                        output_path = os.path.join(output_dir, f"签到统计_{date_str}.docx")
                        self.generate_report(data, output_path)
                        results.append({
                            'source': filename,
                            'output': output_path,
                            'total_expected': data['total_expected'],
                            'total_attended': data['total_attended'],
                            'rate': (data['total_attended'] / data['total_expected'] * 100) if data['total_expected'] > 0 else 0
                        })
        
        return results

if __name__ == "__main__":
    processor = AttendanceProcessor()
    
    test_dir = os.path.join(os.path.dirname(__file__), 'test_data')
    output_dir = os.path.join(os.path.dirname(__file__), 'attendance_reports')
    
    results = processor.batch_process(test_dir, output_dir)
    print(f"处理完成，共处理 {len(results)} 份签到表")
    for r in results:
        print(f"  {r['source']} -> 应到{r['total_expected']}人, 实到{r['total_attended']}人, 签到率{r['rate']:.1f}%")