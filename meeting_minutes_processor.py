import os
import re
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MeetingMinutesProcessor:
    def __init__(self):
        self.minutes_data = []
    
    def parse_minutes(self, file_path):
        if not os.path.exists(file_path):
            return None
        
        _, ext = os.path.splitext(file_path)
        
        if ext.lower() == '.docx':
            return self._parse_docx(file_path)
        elif ext.lower() == '.txt':
            return self._parse_txt(file_path)
        else:
            return None
    
    def _parse_docx(self, file_path):
        doc = Document(file_path)
        content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content.append(text)
        
        return self._extract_info(content)
    
    def _parse_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            content = [line.strip() for line in f if line.strip()]
        
        return self._extract_info(content)
    
    def _extract_info(self, content):
        result = {
            'title': '',
            'date': '',
            'time': '',
            'location': '',
            'participants': [],
            'absentees': [],
            'agenda': [],
            'decisions': [],
            'action_items': [],
            'next_meeting': '',
            'summary': ''
        }
        
        current_section = None
        
        for line in content:
            if re.match(r'会议主题|会议名称|主题', line):
                result['title'] = re.sub(r'(会议主题|会议名称|主题)\s*[：:]', '', line).strip()
            elif re.match(r'时间|会议时间', line):
                time_match = re.search(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日号]?)\s*(\d{1,2}:\d{2}.*)?', line)
                if time_match:
                    result['date'] = time_match.group(1)
                    result['time'] = time_match.group(2) if time_match.group(2) else ''
            elif re.match(r'地点|会议地点', line):
                result['location'] = re.sub(r'(地点|会议地点)\s*[：:]', '', line).strip()
            elif re.match(r'参会人员|出席人员|参会名单|参加人员', line):
                current_section = 'participants'
            elif re.match(r'缺席人员|请假人员', line):
                current_section = 'absentees'
            elif re.match(r'会议议程|议程', line):
                current_section = 'agenda'
            elif re.match(r'会议决议|决议事项|决议', line):
                current_section = 'decisions'
            elif re.match(r'行动项|待办事项|任务安排', line):
                current_section = 'action_items'
            elif re.match(r'下次会议|下次例会', line):
                result['next_meeting'] = re.sub(r'(下次会议|下次例会)\s*[：:]', '', line).strip()
            elif re.match(r'会议总结|总结', line):
                current_section = 'summary'
            elif line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or \
                 line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                item_text = re.sub(r'^[一二三四五六七八九十]+[、.]\s*|\d+[、.]\s*', '', line)
                if current_section == 'participants':
                    result['participants'].extend([p.strip() for p in item_text.split('、') if p.strip()])
                elif current_section == 'absentees':
                    result['absentees'].extend([p.strip() for p in item_text.split('、') if p.strip()])
                elif current_section == 'agenda':
                    result['agenda'].append(item_text)
                elif current_section == 'decisions':
                    result['decisions'].append(item_text)
                elif current_section == 'action_items':
                    result['action_items'].append(item_text)
                elif current_section == 'summary':
                    result['summary'] += item_text + '\n'
            elif current_section == 'participants':
                result['participants'].extend([p.strip() for p in line.split('、') if p.strip()])
            elif current_section == 'absentees':
                result['absentees'].extend([p.strip() for p in line.split('、') if p.strip()])
        
        if not result['date']:
            result['date'] = datetime.now().strftime('%Y年%m月%d日')
        
        return result
    
    def generate_structured_minutes(self, data, output_path):
        doc = Document()
        
        style = doc.styles['Heading 1']
        style.font.size = Pt(16)
        style.font.bold = True
        
        doc.add_heading(data['title'] or '会议纪要', level=1)
        
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = '会议时间'
        info_table.cell(0, 1).text = f"{data['date']} {data['time']}"
        info_table.cell(1, 0).text = '会议地点'
        info_table.cell(1, 1).text = data['location'] or '未指定'
        info_table.cell(2, 0).text = '参会人数'
        info_table.cell(2, 1).text = f"{len(data['participants'])}人"
        
        if data['participants']:
            doc.add_heading('参会人员', level=2)
            doc.add_paragraph('、'.join(data['participants']))
        
        if data['absentees']:
            doc.add_heading('缺席人员', level=2)
            doc.add_paragraph('、'.join(data['absentees']))
        
        if data['agenda']:
            doc.add_heading('会议议程', level=2)
            for i, item in enumerate(data['agenda'], 1):
                doc.add_paragraph(f"{i}. {item}", style='List Number')
        
        if data['decisions']:
            doc.add_heading('会议决议', level=2)
            for i, item in enumerate(data['decisions'], 1):
                doc.add_paragraph(f"{i}. {item}", style='List Number')
        
        if data['action_items']:
            doc.add_heading('行动项', level=2)
            for i, item in enumerate(data['action_items'], 1):
                doc.add_paragraph(f"{i}. {item}", style='List Number')
        
        if data['next_meeting']:
            doc.add_heading('下次会议安排', level=2)
            doc.add_paragraph(data['next_meeting'])
        
        if data['summary']:
            doc.add_heading('会议总结', level=2)
            doc.add_paragraph(data['summary'])
        
        doc.save(output_path)
        return output_path
    
    def auto_archive(self, source_dir, archive_base_dir):
        os.makedirs(archive_base_dir, exist_ok=True)
        
        for filename in os.listdir(source_dir):
            if filename.endswith('.docx') or filename.endswith('.txt'):
                if '会议纪要' in filename or '纪要' in filename:
                    source_path = os.path.join(source_dir, filename)
                    data = self.parse_minutes(source_path)
                    
                    if data:
                        date_str = data['date'].replace('/', '-').replace('年', '-').replace('月', '-').replace('日', '').replace('号', '')
                        date_str = re.sub(r'[-]+', '-', date_str).strip('-')
                        
                        month_dir = os.path.join(archive_base_dir, date_str[:7] if len(date_str) >= 7 else datetime.now().strftime('%Y-%m'))
                        os.makedirs(month_dir, exist_ok=True)
                        
                        new_filename = f"会议纪要_{date_str}.docx"
                        output_path = os.path.join(month_dir, new_filename)
                        
                        self.generate_structured_minutes(data, output_path)
                        print(f"已归档: {filename} -> {output_path}")
    
    def batch_process(self, input_dir, output_dir):
        os.makedirs(output_dir, exist_ok=True)
        results = []
        
        for filename in os.listdir(input_dir):
            if filename.endswith('.docx') or filename.endswith('.txt'):
                if '会议纪要' in filename or '纪要' in filename:
                    source_path = os.path.join(input_dir, filename)
                    data = self.parse_minutes(source_path)
                    
                    if data:
                        date_str = data['date'].replace('/', '-').replace('年', '-').replace('月', '-').replace('日', '').replace('号', '')
                        date_str = re.sub(r'[-]+', '-', date_str).strip('-')
                        
                        output_path = os.path.join(output_dir, f"会议纪要_{date_str}.docx")
                        self.generate_structured_minutes(data, output_path)
                        results.append({
                            'source': filename,
                            'output': output_path,
                            'participants_count': len(data['participants']),
                            'action_items_count': len(data['action_items'])
                        })
        
        return results

if __name__ == "__main__":
    processor = MeetingMinutesProcessor()
    
    test_dir = os.path.join(os.path.dirname(__file__), 'test_data')
    output_dir = os.path.join(os.path.dirname(__file__), 'processed_minutes')
    
    results = processor.batch_process(test_dir, output_dir)
    print(f"处理完成，共处理 {len(results)} 份会议纪要")
    for r in results:
        print(f"  {r['source']} -> {r['output']}")