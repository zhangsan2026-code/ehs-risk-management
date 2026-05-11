import os
import re
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

class TrainingRecordsProcessor:
    def __init__(self):
        self.training_records = []
    
    def parse_training_record(self, file_path):
        if not os.path.exists(file_path):
            return None
        
        _, ext = os.path.splitext(file_path)
        
        if ext.lower() == '.xlsx' or ext.lower() == '.xls':
            return self._parse_excel(file_path)
        elif ext.lower() == '.docx':
            return self._parse_docx(file_path)
        else:
            return None
    
    def _parse_excel(self, file_path):
        try:
            df = pd.read_excel(file_path)
        except Exception as e:
            print(f"读取Excel文件失败: {e}")
            return None
        
        result = {
            'title': '',
            'date': '',
            'trainer': '',
            'department': '',
            'total_participants': 0,
            'participants': [],
            'evaluation_scores': [],
            'overall_score': 0,
            'comments': [],
            'effect_summary': ''
        }
        
        for col in df.columns:
            col_str = str(col)
            
            if '培训名称' in col_str or '课程名称' in col_str:
                result['title'] = df[col].dropna().iloc[0] if not df[col].dropna().empty else ''
            elif '日期' in col_str:
                for val in df[col].dropna():
                    date_match = re.search(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日号]?)', str(val))
                    if date_match:
                        result['date'] = date_match.group(1)
                        break
            elif '培训师' in col_str or '讲师' in col_str:
                result['trainer'] = df[col].dropna().iloc[0] if not df[col].dropna().empty else ''
            elif '部门' in col_str:
                result['department'] = df[col].dropna().iloc[0] if not df[col].dropna().empty else ''
        
        name_col = None
        score_col = None
        comment_col = None
        
        for col in df.columns:
            col_str = str(col)
            if any(keyword in col_str for keyword in ['姓名', '人员', '学员']):
                name_col = col
            elif any(keyword in col_str for keyword in ['评分', '成绩', '分数', '评价']):
                score_col = col
            elif any(keyword in col_str for keyword in ['心得', '反馈', '建议', '评语']):
                comment_col = col
        
        if name_col:
            for _, row in df.iterrows():
                name = str(row[name_col]).strip()
                if not name or name in ['姓名', '人员', '学员', 'nan']:
                    continue
                
                score = None
                if score_col:
                    try:
                        score = float(row[score_col])
                        result['evaluation_scores'].append(score)
                    except:
                        pass
                
                comment = ''
                if comment_col:
                    comment = str(row[comment_col]).strip()
                    if comment and comment != 'nan':
                        result['comments'].append({'name': name, 'comment': comment})
                
                result['participants'].append({
                    'name': name,
                    'score': score,
                    'comment': comment
                })
        
        result['total_participants'] = len(result['participants'])
        
        if result['evaluation_scores']:
            result['overall_score'] = sum(result['evaluation_scores']) / len(result['evaluation_scores'])
        
        return result
    
    def _parse_docx(self, file_path):
        doc = Document(file_path)
        result = {
            'title': '',
            'date': '',
            'trainer': '',
            'department': '',
            'total_participants': 0,
            'participants': [],
            'evaluation_scores': [],
            'overall_score': 0,
            'comments': [],
            'effect_summary': ''
        }
        
        content = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content.append(text)
        
        for line in content:
            if '培训名称' in line or '课程名称' in line:
                result['title'] = re.sub(r'(培训名称|课程名称)\s*[：:]', '', line).strip()
            elif '日期' in line:
                date_match = re.search(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日号]?)', line)
                if date_match:
                    result['date'] = date_match.group(1)
            elif '培训师' in line or '讲师' in line:
                result['trainer'] = re.sub(r'(培训师|讲师)\s*[：:]', '', line).strip()
            elif '部门' in line:
                result['department'] = re.sub(r'部门\s*[：:]', '', line).strip()
            elif '得分' in line or '评分' in line:
                score_match = re.search(r'(\d+\.?\d*)', line)
                if score_match:
                    try:
                        result['evaluation_scores'].append(float(score_match.group(1)))
                    except:
                        pass
        
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                
                if len(row_data) >= 1:
                    name = row_data[0]
                    if name and name not in ['姓名', '人员', 'nan']:
                        score = None
                        if len(row_data) >= 2:
                            try:
                                score = float(row_data[1])
                                result['evaluation_scores'].append(score)
                            except:
                                pass
                        
                        comment = row_data[2] if len(row_data) >= 3 else ''
                        if comment:
                            result['comments'].append({'name': name, 'comment': comment})
                        
                        result['participants'].append({
                            'name': name,
                            'score': score,
                            'comment': comment
                        })
        
        result['total_participants'] = len(result['participants'])
        
        if result['evaluation_scores']:
            result['overall_score'] = sum(result['evaluation_scores']) / len(result['evaluation_scores'])
        
        return result
    
    def generate_training_report(self, data, output_path):
        doc = Document()
        
        style = doc.styles['Heading 1']
        style.font.size = Pt(16)
        style.font.bold = True
        
        doc.add_heading(f"{data['title'] or '培训'}效果评估报告", level=1)
        
        info_table = doc.add_table(rows=2, cols=3)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = '培训日期'
        info_table.cell(0, 1).text = data['date'] or datetime.now().strftime('%Y年%m月%d日')
        info_table.cell(0, 2).text = '培训师'
        info_table.cell(1, 0).text = data['trainer'] or '未指定'
        info_table.cell(1, 1).text = '部门'
        info_table.cell(1, 2).text = data['department'] or '未指定'
        
        for row in info_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        stats_table.cell(0, 0).text = '参与人数'
        stats_table.cell(0, 1).text = str(data['total_participants'])
        
        doc.add_paragraph(f"\n平均评分: {data['overall_score']:.1f}分")
        
        if data['participants']:
            doc.add_heading('学员名单及评分', level=2)
            participants_table = doc.add_table(rows=1, cols=3)
            participants_table.style = 'Table Grid'
            participants_table.cell(0, 0).text = '序号'
            participants_table.cell(0, 1).text = '姓名'
            participants_table.cell(0, 2).text = '评分'
            
            for i, participant in enumerate(data['participants'], 1):
                row = participants_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = participant['name']
                row.cells[2].text = str(participant['score']) if participant['score'] else '-'
        
        if data['comments']:
            doc.add_heading('学员反馈', level=2)
            for i, comment in enumerate(data['comments'], 1):
                doc.add_paragraph(f"{i}. {comment['name']}: {comment['comment']}")
        
        if data['effect_summary']:
            doc.add_heading('培训效果总结', level=2)
            doc.add_paragraph(data['effect_summary'])
        
        doc.save(output_path)
        return output_path
    
    def batch_process(self, input_dir, output_dir):
        os.makedirs(output_dir, exist_ok=True)
        results = []
        
        for filename in os.listdir(input_dir):
            if filename.endswith('.xlsx') or filename.endswith('.xls') or filename.endswith('.docx'):
                if '培训' in filename or '效果' in filename:
                    source_path = os.path.join(input_dir, filename)
                    data = self.parse_training_record(source_path)
                    
                    if data:
                        date_str = data['date'].replace('/', '-').replace('年', '-').replace('月', '-').replace('日', '').replace('号', '')
                        date_str = re.sub(r'[-]+', '-', date_str).strip('-')
                        if len(date_str) < 7:
                            date_str = datetime.now().strftime('%Y-%m')
                        
                        output_path = os.path.join(output_dir, f"培训效果报告_{date_str}.docx")
                        self.generate_training_report(data, output_path)
                        results.append({
                            'source': filename,
                            'output': output_path,
                            'participants': data['total_participants'],
                            'avg_score': data['overall_score']
                        })
        
        return results

if __name__ == "__main__":
    processor = TrainingRecordsProcessor()
    
    test_dir = os.path.join(os.path.dirname(__file__), 'test_data')
    output_dir = os.path.join(os.path.dirname(__file__), 'training_reports')
    
    results = processor.batch_process(test_dir, output_dir)
    print(f"处理完成，共处理 {len(results)} 份培训记录")
    for r in results:
        print(f"  {r['source']} -> 参与{r['participants']}人, 平均分{r['avg_score']:.1f}")